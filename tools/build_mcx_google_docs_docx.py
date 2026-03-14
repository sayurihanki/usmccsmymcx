from __future__ import annotations

import io
from pathlib import Path
from typing import Any
from urllib.request import urlopen

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "output" / "doc"
OUTPUT_PATH = OUTPUT_DIR / "mcx-google-docs-blocks.docx"


def image_cell(url: str, alt: str, width: float = 1.45) -> dict[str, Any]:
    return {
        "type": "image",
        "url": url,
        "alt": alt,
        "width": width,
    }


def link_cell(text: str, url: str) -> dict[str, Any]:
    return {
        "type": "link",
        "text": text,
        "url": url,
    }


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def add_hyperlink(paragraph, text: str, url: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "C1121F")
    r_pr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)

    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)

    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def clear_cell(cell) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if paragraph.runs:
        for run in paragraph.runs:
            run.clear()


def write_cell(cell, value: Any) -> None:
    clear_cell(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

    if isinstance(value, dict) and value.get("type") == "image":
        try:
            with urlopen(value["url"]) as response:
                image_bytes = io.BytesIO(response.read())
            run = paragraph.add_run()
            run.add_picture(image_bytes, width=Inches(value.get("width", 1.45)))
            if value.get("alt"):
                caption = cell.add_paragraph(value["alt"])
                caption.style = "Body Text"
                caption.runs[0].italic = True
                caption.runs[0].font.size = Pt(8)
                caption.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        except Exception:
            fallback = paragraph.add_run(f"[Image] {value.get('alt', 'Image')}")
            fallback.font.bold = True
            fallback.font.size = Pt(10)
            source = cell.add_paragraph(value["url"])
            source.runs[0].font.size = Pt(8)
            source.runs[0].font.color.rgb = RGBColor(0x6B, 0x72, 0x80)
        return

    if isinstance(value, dict) and value.get("type") == "link":
        add_hyperlink(paragraph, value["text"], value["url"])
        return

    text = str(value or "")
    for index, line in enumerate(text.replace("<br>", "\n").split("\n")):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(line)
        run.font.size = Pt(10.5)


def add_table(document: Document, block_name: str, columns: int, rows: list[list[Any]], widths: list[int]) -> None:
    table = document.add_table(rows=len(rows) + 1, cols=columns)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    header = table.rows[0].cells
    merged = header[0]
    for cell in header[1:]:
        merged = merged.merge(cell)
    merged.text = block_name
    set_cell_shading(merged, "C1121F")
    header_run = merged.paragraphs[0].runs[0]
    header_run.font.bold = True
    header_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    header_run.font.size = Pt(11)

    for row_index, row_values in enumerate(rows, start=1):
        row = table.rows[row_index]
        padded = row_values + [""] * (columns - len(row_values))
        for col_index, value in enumerate(padded):
            cell = row.cells[col_index]
            set_cell_width(cell, widths[col_index])
            write_cell(cell, value)

    document.add_paragraph("")


def add_heading(document: Document, text: str, level: int = 2) -> None:
    paragraph = document.add_paragraph()
    paragraph.style = f"Heading {level}"
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x12, 0x20, 0x33)


def add_note(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x5E, 0x6F, 0x82)


SECTIONS: list[dict[str, Any]] = [
    {
        "heading": "1. Homepage Metadata",
        "note": "Paste this table into the homepage Google Doc metadata area.",
        "block": "metadata",
        "columns": 2,
        "widths": [2400, 6960],
        "rows": [
            ["template", "mcx-home"],
            ["theme", "mcx"],
            ["nav", "/fragments/mcx-nav"],
            ["footer", "/fragments/mcx-footer"],
        ],
    },
    {
        "heading": "2. Header Fragment",
        "note": "Paste this into a dedicated fragment doc for /fragments/mcx-nav.",
        "block": "mcx-nav-data",
        "columns": 4,
        "widths": [1600, 2500, 2500, 2760],
        "rows": [
            ["config", "", "search-placeholder", "Search products, brands, gear..."],
            ["config", "", "sign-in-label", "Sign In"],
            ["config", "", "sign-in-url", "/account"],
            ["hot", "", "Deals & Offers", "/deals"],
            ["item", "", "Shop All", "/shop"],
            ["item", "", "Apparel & Footwear", "/apparel-footwear"],
            ["group", "Apparel & Footwear", "Men's", ""],
            ["group", "Apparel & Footwear", "Women's", ""],
            ["group", "Apparel & Footwear", "Footwear", ""],
            ["link", "Apparel & Footwear > Men's", "Uniforms & Cammies", "/mens/uniforms"],
            ["link", "Apparel & Footwear > Men's", "T-Shirts & Polos", "/mens/tops"],
            ["link", "Apparel & Footwear > Men's", "Athletic Wear", "/mens/athletics"],
            ["link", "Apparel & Footwear > Women's", "Uniforms", "/womens/uniforms"],
            ["link", "Apparel & Footwear > Women's", "Activewear", "/womens/activewear"],
            ["link", "Apparel & Footwear > Women's", "Accessories", "/womens/accessories"],
            ["link", "Apparel & Footwear > Footwear", "Tactical Boots", "/footwear/boots"],
            ["link", "Apparel & Footwear > Footwear", "Athletic Shoes", "/footwear/athletic"],
            ["link", "Apparel & Footwear > Footwear", "Work & Safety", "/footwear/work"],
            ["feature", "Apparel & Footwear", "Clearance Apparel", "/clearance/apparel"],
            ["feature", "Apparel & Footwear", "Top Rated Boots", "/boots"],
            ["feature", "Apparel & Footwear", "New Arrivals", "/new"],
            ["item", "", "Electronics", "/electronics"],
            ["group", "Electronics", "Computers", ""],
            ["group", "Electronics", "Mobile & Audio", ""],
            ["group", "Electronics", "Gaming", ""],
            ["link", "Electronics > Computers", "MacBooks & Laptops", "/electronics/computers"],
            ["link", "Electronics > Computers", "Tablets & iPads", "/electronics/tablets"],
            ["link", "Electronics > Mobile & Audio", "Smartphones", "/electronics/mobile"],
            ["link", "Electronics > Mobile & Audio", "Headphones", "/electronics/audio"],
            ["link", "Electronics > Gaming", "Consoles", "/electronics/gaming"],
            ["link", "Electronics > Gaming", "Games & Software", "/electronics/gaming"],
            ["feature", "Electronics", "Apple Products", "/electronics/apple"],
            ["feature", "Electronics", "Gaming Deals", "/electronics/gaming"],
            ["item", "", "Home & Living", "/home"],
            ["group", "Home & Living", "Furniture", ""],
            ["group", "Home & Living", "Kitchen", ""],
            ["group", "Home & Living", "Bedding & Bath", ""],
            ["link", "Home & Living > Furniture", "Bedroom", "/home/furniture"],
            ["link", "Home & Living > Kitchen", "Appliances", "/home/kitchen"],
            ["link", "Home & Living > Bedding & Bath", "Towels", "/home/bath"],
            ["item", "", "Sports & Outdoors", "/sports-outdoors"],
            ["item", "", "Uniforms", "/uniforms"],
            ["item", "", "Kids & Toys", "/kids-toys"],
        ],
    },
    {
        "heading": "3. Footer Fragment",
        "note": "Paste this into a dedicated fragment doc for /fragments/mcx-footer.",
        "block": "mcx-footer-data",
        "columns": 4,
        "widths": [1600, 2000, 2500, 3260],
        "rows": [
            ["config", "brand", "about", "Serving Marines and their families since 1897. Tax-free shopping that returns 100% of profits to the Marine Corps community through MCCS programs and activities."],
            ["config", "site", "note", "© 2026 Marine Corps Exchange (MCX) / MCCS. All rights reserved. Tax-free shopping for authorized DoD patrons."],
            ["config", "site", "flag", "Proudly Serving Since 1897"],
            ["social", "facebook", "Facebook", "https://www.facebook.com/"],
            ["social", "instagram", "Instagram", "https://www.instagram.com/"],
            ["social", "twitter", "Twitter", "https://www.x.com/"],
            ["social", "tiktok", "TikTok", "https://www.tiktok.com/"],
            ["column", "Shop", "New Arrivals", "/new"],
            ["column", "Shop", "Men's", "/mens"],
            ["column", "Shop", "Women's", "/womens"],
            ["column", "Shop", "Kids & Baby", "/kids"],
            ["column", "Shop", "Electronics", "/electronics"],
            ["column", "Shop", "Home & Garden", "/home"],
            ["column", "Shop", "Sports & Outdoors", "/sports-outdoors"],
            ["column", "Shop", "Sale & Clearance", "/clearance"],
            ["column", "Services", "Find a Store", "/stores"],
            ["column", "Services", "Ship to Store", "/shipping"],
            ["column", "Services", "Returns & Exchanges", "/returns"],
            ["column", "Services", "Order Status", "/orders"],
            ["column", "Services", "Military Layaway", "/services"],
            ["column", "Services", "Gift Cards", "/gift-cards"],
            ["column", "Services", "Price Match", "/price-match"],
            ["column", "About MCX", "Our Mission", "/about"],
            ["column", "About MCX", "MCCS Programs", "/mccs"],
            ["column", "About MCX", "Patron Eligibility", "/eligibility"],
            ["column", "About MCX", "Vendor Partnerships", "/partners"],
            ["column", "About MCX", "Careers at MCX", "/careers"],
            ["column", "About MCX", "Newsroom", "/newsroom"],
            ["column", "Support", "Help Center", "/help"],
            ["column", "Support", "Contact Us", "/contact"],
            ["column", "Support", "Feedback", "/feedback"],
            ["column", "Support", "Accessibility", "/accessibility"],
            ["column", "Support", "Security Center", "/security"],
            ["legal", "", "Privacy Policy", "/privacy"],
            ["legal", "", "Terms of Use", "/terms"],
            ["legal", "", "Cookie Policy", "/cookies"],
            ["legal", "", "Accessibility", "/accessibility"],
        ],
    },
    {
        "heading": "4. Announcement Bar",
        "note": "Paste on the page body.",
        "block": "mcx-announcement-bar",
        "columns": 2,
        "widths": [2400, 6960],
        "rows": [
            ["dismissible", "true"],
            ["Free Shipping $50+", ""],
            ["Tax-Free Shopping - Authorized Patrons Only", ""],
            ["Spring 2026 Collection - Now Live", ""],
            ["Ship to Store: 3-5 Days Stateside", ""],
            ["100% of Profits Fund Marine Programs", ""],
        ],
    },
    {
        "heading": "5. Hero",
        "note": "Explicit hero block. This is not auto-generated.",
        "block": "mcx-hero",
        "columns": 2,
        "widths": [2400, 6960],
        "rows": [
            ["eyebrow", "Spring Collection - 2026 - Tax-Free"],
            ["heading-line-1", "OUTFITTED"],
            ["heading-line-2", "FOR THE"],
            ["heading-line-3", "mission & beyond"],
            ["description", "Serving Marines and their families since 1897. Premium brands, exclusive savings, and tax-free shopping - exclusively for those who serve."],
            ["image", image_cell("https://images.unsplash.com/photo-1519415943484-9fa1873496d4?w=1200&q=80", "Marine Corps collection hero", 1.9)],
            ["primary-cta", link_cell("Shop Now", "#products")],
            ["secondary-cta", link_cell("View Deals", "#deals")],
            ["status-badge-1", "SYS: MCX-2026"],
            ["status-badge-2", "STATUS: ACTIVE"],
            ["status-badge-3", "PATRON: AUTHORIZED"],
            ["stat-1-value", "20%+"],
            ["stat-1-label", "Average Savings"],
            ["stat-2-value", "33M+"],
            ["stat-2-label", "Yearly Transactions"],
            ["stat-3-value", "Tax Free"],
            ["stat-3-label", "Exclusive Benefit"],
            ["stat-4-value", "127+"],
            ["stat-4-label", "Store Locations"],
        ],
    },
    {
        "heading": "6. Ticker",
        "note": "Single-column ticker rows.",
        "block": "mcx-ticker",
        "columns": 1,
        "widths": [9360],
        "rows": [
            ["Free Shipping|on $50+"],
            ["Tax-Free|Shopping"],
            ["Under Armour|featured this week"],
            ["New Arrivals|every week"],
            ["Apple|products available"],
            ["Save 20%+|vs retail"],
        ],
    },
    {
        "heading": "7. Benefits",
        "note": "Three-column icon + title + subtitle rows.",
        "block": "mcx-benefits",
        "columns": 3,
        "widths": [1900, 2800, 4660],
        "rows": [
            ["taxfree", "Tax-Free Shopping", "Exclusive benefit for authorized patrons"],
            ["shipping", "Free Shipping $50+", "Ship to store in 3-5 days stateside"],
            ["quality", "Quality Guaranteed", "Premium brands below market price"],
            ["givesback", "100% Gives Back", "All profits support Marine community"],
        ],
    },
    {
        "heading": "8. Category Section Metadata",
        "note": "Paste immediately above the category grid block.",
        "block": "section-metadata",
        "columns": 2,
        "widths": [2400, 6960],
        "rows": [["style", "sec, sec-mid"]],
    },
    {
        "heading": "9. Category Grid",
        "note": "Four-column category rows.",
        "block": "mcx-category-grid",
        "columns": 4,
        "widths": [1300, 2300, 2300, 3460],
        "rows": [
            ["label", "Browse the Exchange", "", ""],
            ["title", "SHOP YOUR WAY", "", ""],
            ["\\U0001f455", "Apparel", "2,400+ items", link_cell("Shop Apparel", "/apparel")],
            ["\\U0001f45f", "Footwear", "890+ styles", link_cell("Shop Footwear", "/footwear")],
            ["\\U0001f4bb", "Electronics", "1,200+ items", link_cell("Shop Electronics", "/electronics")],
            ["\\U0001f3e0", "Home & Living", "3,100+ items", link_cell("Shop Home", "/home")],
            ["\\U0001f3cb", "Sports", "760+ items", link_cell("Shop Sports", "/sports-outdoors")],
            ["\\U0001f484", "Beauty", "940+ items", link_cell("Shop Beauty", "/beauty")],
            ["\\U0001f3ae", "Gaming", "520+ items", link_cell("Shop Gaming", "/gaming")],
            ["\\U0001f392", "Tactical Gear", "380+ items", link_cell("Shop Tactical", "/tactical")],
            ["\\U0001f9f8", "Kids & Baby", "1,500+ items", link_cell("Shop Kids", "/kids")],
            ["\\u26fa", "Outdoors", "670+ items", link_cell("Shop Outdoors", "/sports-outdoors")],
            ["\\U0001f43e", "Pets", "290+ items", link_cell("Shop Pets", "/pets")],
            ["\\U0001f381", "Gifts", "800+ ideas", link_cell("Shop Gifts", "/gifts")],
        ],
    },
    {
        "heading": "10. Product Section Metadata",
        "note": "Paste immediately above the product cards block.",
        "block": "section-metadata",
        "columns": 2,
        "widths": [2400, 6960],
        "rows": [["style", "sec, sec-dark"]],
    },
    {
        "heading": "11. Product Cards",
        "note": "Four-column product rows with image, name, metrics, and config.",
        "block": "mcx-product-cards",
        "columns": 4,
        "widths": [1700, 2500, 2300, 2860],
        "rows": [
            ["label", "Curated For You", "", ""],
            ["title", "NEW ARRIVALS", "", ""],
            ["view-link", link_cell("View All Products", "/products"), "", ""],
            ["tabs", "All Items|all, Apparel|apparel, Electronics|electronics, Footwear|footwear, Tactical|tactical, Sports|sports", "", ""],
            ["default-tab", "all", "", ""],
            [image_cell("https://images.unsplash.com/photo-1542291026-7eec264c27ff?w=500&q=75", "Tactical boot", 1.15), "Danner\nTachyon 8 Tactical Boot GTX", "rating: 5\nreviews: 248\nprice: 189.99\noriginal: 239.99\nchip: Save $50", "category: footwear\nflags: New\nemoji: \\U0001f97e\nsizes: 8, 9, 10, 11, 12"],
            [image_cell("https://images.unsplash.com/photo-1505740420928-5e560c06d30e?w=500&q=75", "Sony headphones", 1.15), "Sony\nWH-1000XM5 Noise Canceling Headphones", "rating: 5\nreviews: 1024\nprice: 279.99\noriginal: 349.99\nchip: Save $70", "category: electronics\nflags: Sale\nemoji: \\U0001f3a7"],
            [image_cell("https://images.unsplash.com/photo-1556821840-3a63f15732ce?w=500&q=75", "Under Armour vest", 1.15), "Under Armour\nUA Storm ColdGear Reactor Vest", "rating: 4\nreviews: 187\nprice: 89.99\noriginal: 120.00\nchip: Save 25%", "category: apparel\nflags: MCX Exclusive, New\nemoji: \\U0001f9e5\nsizes: XS, S, M, L, XL"],
            [image_cell("https://images.unsplash.com/photo-1523275335684-37898b6baf30?w=500&q=75", "Garmin watch", 1.15), "Garmin\nInstinct 2X Solar Tactical Edition", "rating: 5\nreviews: 512\nprice: 424.99\noriginal: 499.99\nchip: Save $75", "category: electronics\nflags: 15% Off\nemoji: \\u231a"],
            [image_cell("https://images.unsplash.com/photo-1514989940723-e8e51635b782?w=500&q=75", "Brooks running shoes", 1.15), "Brooks\nGhost 16 Road Running Shoes", "rating: 5\nreviews: 893\nprice: 119.99\noriginal: 139.99\nchip: Save $20", "category: footwear\nflags: New\nemoji: \\U0001f45f\nsizes: 8, 9, 10, 11, 12"],
            [image_cell("https://images.unsplash.com/photo-1544244015-0df4b3ffc6b0?w=500&q=75", "iPad Air", 1.15), "Apple\niPad Air M2 11-inch 128GB WiFi", "rating: 5\nreviews: 2104\nprice: 549.99\noriginal: 599.00\nchip: Save $49", "category: electronics\nflags: MCX Bundle\nemoji: \\U0001f4f1"],
            [image_cell("https://images.unsplash.com/photo-1553062407-98eeb64c6a62?w=500&q=75", "5.11 backpack", 1.15), "5.11 Tactical\nRush 72 2.0 Backpack 55L", "rating: 5\nreviews: 631\nprice: 154.99\noriginal: 185.00\nchip: Save $30", "category: tactical\nflags: Top Rated\nemoji: \\U0001f392"],
            [image_cell("https://images.unsplash.com/photo-1486401899868-0e435ed85128?w=500&q=75", "PlayStation 5", 1.15), "Sony PlayStation\nPlayStation 5 Console Bundle", "rating: 5\nreviews: 3441\nprice: 449.99\noriginal: 499.99\nchip: Save $50", "category: electronics\nflags: In Stock\nemoji: \\U0001f3ae"],
        ],
    },
    {
        "heading": "12. Deal Countdown",
        "note": "Two-column key/value block.",
        "block": "mcx-deal-countdown",
        "columns": 2,
        "widths": [2400, 6960],
        "rows": [
            ["label", "Flash Sale - Limited Time"],
            ["title", "SPRING CLEARANCE UP TO 40% OFF"],
            ["description", "Exclusive deals across apparel, electronics, tactical gear, and more. Tax-free pricing for all authorized MCX patrons - active duty, veterans, retirees, and families."],
            ["end-datetime", "2026-04-01T23:59:59-07:00"],
            ["cta", link_cell("Shop All Deals", "/deals")],
            ["cta-note", "Do not miss out - deals expire when the timer hits zero."],
            ["ended-text", "Ended"],
            ["hide-cta-when-ended", "false"],
        ],
    },
    {
        "heading": "13. Promo Strip",
        "note": "Two-column key/value block.",
        "block": "mcx-promo-strip",
        "columns": 2,
        "widths": [2400, 6960],
        "rows": [
            ["badge", "Spring 2026 - Limited Time"],
            ["title", "EARN MORE, SPEND LESS"],
            ["description", "Every purchase supports Marine Corps Community Services. You have earned these benefits - shop smarter, save more, and give back to the community that has your back."],
            ["cta", link_cell("Shop Clearance", "/clearance")],
        ],
    },
    {
        "heading": "14. Featured Collections Section Metadata",
        "note": "Paste immediately above the featured collections block.",
        "block": "section-metadata",
        "columns": 2,
        "widths": [2400, 6960],
        "rows": [["style", "sec, sec-dark"]],
    },
    {
        "heading": "15. Featured Collections",
        "note": "Four-column image + tag + title + CTA rows.",
        "block": "mcx-featured-collections",
        "columns": 4,
        "widths": [1700, 2300, 2800, 2560],
        "rows": [
            ["label", "Curated Collections", "", ""],
            ["title", "SHOP THE LOOK", "", ""],
            ["view-link", link_cell("All Collections", "/collections"), "", ""],
            [image_cell("https://images.unsplash.com/photo-1517836357463-d25dfeac3438?w=800&q=80", "Tactical collection", 1.2), "Spring 2026 - Men's", "TACTICAL COLLECTION", link_cell("Shop Men's", "/collections/mens")],
            [image_cell("https://images.unsplash.com/photo-1571019614242-c5c5dee9f50b?w=600&q=80", "Performance training", 1.2), "Athletics & Training", "PERFORMANCE TRAINING", link_cell("Shop Gear", "/collections/training")],
            [image_cell("https://images.unsplash.com/photo-1555041469-a586c61ea9bc?w=600&q=80", "Home essentials", 1.2), "Home & Living", "FRESH HOME ESSENTIALS", link_cell("Shop Home", "/collections/home")],
        ],
    },
    {
        "heading": "16. Brands",
        "note": "Single-column brand rows.",
        "block": "mcx-brands",
        "columns": 1,
        "widths": [9360],
        "rows": [
            ["Apple"],
            ["Samsung"],
            ["Under Armour"],
            ["Brooks"],
            ["5.11 Tactical"],
            ["Garmin"],
            ["Danner"],
        ],
    },
    {
        "heading": "17. Editorial Cards",
        "note": "Four-column editorial rows.",
        "block": "mcx-editorial-cards",
        "columns": 4,
        "widths": [1700, 1000, 1800, 4860],
        "rows": [
            [image_cell("https://images.unsplash.com/photo-1565462905350-4a1b2a4b3a8a?w=600&q=80", "Uniform guide", 1.15), "01", "Uniform Guide", "Everything You Need for Inspection-Ready Dress Blues\nA complete guide to proper uniform wear, care, and every authorized item at MCX - from ribbons to shoes."],
            [image_cell("https://images.unsplash.com/photo-1571019613454-1cb2f99b2d8b?w=600&q=80", "Fitness gear", 1.15), "02", "Fitness", "Top PT Gear to Crush Your PFT Score This Year\nOur experts selected the best running shoes and training gear built for the demands of Marine physical training."],
            [image_cell("https://images.unsplash.com/photo-1556909114-f6e7ad7d3136?w=600&q=80", "Family home", 1.15), "03", "Family Living", "Set Up Your New BAH Home Without Breaking the Bank\nSmart buys for furnishing and equipping your first home - all tax-free and ready to ship to your nearest MCX."],
        ],
    },
    {
        "heading": "18. Newsletter",
        "note": "Final homepage block.",
        "block": "mcx-newsletter",
        "columns": 2,
        "widths": [2400, 6960],
        "rows": [
            ["label", "Stay in the Loop"],
            ["title", "GET EXCLUSIVE DEALS"],
            ["description", "Early access to MCX deals, new arrivals, and Marine community news. Reserved for authorized patrons - because you have earned these benefits."],
            ["placeholder", "Your military email address..."],
            ["button-text", "Subscribe"],
            ["note", "AUTHORIZED PATRONS ONLY - UNSUBSCRIBE ANYTIME - NO SPAM"],
        ],
    },
]


def build_document() -> Path:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run("MCX Google Docs Block Tables")
    title_run.font.name = "Arial"
    title_run.font.bold = True
    title_run.font.size = Pt(20)

    subtitle = document.add_paragraph()
    subtitle_run = subtitle.add_run(
        "This .docx contains native Word tables for each MCX block and fragment so the file can be imported into Google Docs and copied from there."
    )
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor(0x5E, 0x6F, 0x82)

    for index, section_data in enumerate(SECTIONS):
        add_heading(document, section_data["heading"], level=2)
        add_note(document, section_data["note"])
        add_table(
            document,
            section_data["block"],
            section_data["columns"],
            section_data["rows"],
            section_data["widths"],
        )
        if index in {2, 9, 17} and index != len(SECTIONS) - 1:
            document.add_section(WD_SECTION.NEW_PAGE)

    document.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build_document()
    verify = Document(path)
    print(f"Wrote {path}")
    print(f"Table count: {len(verify.tables)}")
